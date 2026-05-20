import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """
    Sets background color for table cells.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_roadmap():
    base_dir = r"c:\Users\thale\Documents\IFvest\api"
    doc = Document()
    
    # Configure styles (font: Calibri, base size: 11)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B)
    
    # Title Section
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("IFVest – Question Extraction API\nRoteiro de Estudos & Engenharia de Dados")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Premium Dark Blue
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("Roteiro prático para o domínio de extração de dados, validação com Pydantic e OCR")
    sub_run.font.name = 'Segoe UI'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # Section 1: Introduction
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Introdução e Visão Geral")
    h1_run.font.name = 'Segoe UI'
    h1_run.font.size = Pt(16)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    p = doc.add_paragraph(
        "A extração de questões de exames de larga escala (como ENEM, FUVEST e exames modulares como o PISM) é "
        "uma tarefa fundamental para a democratização e digitalização da educação. Este projeto consiste em uma "
        "Engenharia de Dados completa de ponta a ponta: desde a ingestão automatizada de arquivos brutos (PDFs das provas), "
        "passando pela normalização de texturas e encodings Unicode, segmentação inteligente com Expressões Regulares, "
        "validação estrita com Pydantic e fornecimento via REST API moderna com FastAPI."
    )
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    
    p2 = doc.add_paragraph(
        "Durante a fase de pesquisa, realizamos uma descoberta crítica: todas as provas fornecidas possuem "
        "camadas nativas de texto digital selecionável. Isso elimina a necessidade primária de processamento OCR lento e "
        "suscetível a erros, possibilitando uma taxa de acerto de 100% de caracteres e uma velocidade de processamento instantânea. "
        "Para novos cadernos de provas escaneados, realizamos um estudo comparativo de algoritmos de OCR que descrevemos neste roteiro."
    )
    p2.paragraph_format.space_after = Pt(20)
    
    # Section 2: OCR Feasibility Study
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Estudo de Viabilidade: Algoritmos de OCR")
    h2_run.font.name = 'Segoe UI'
    h2_run.font.size = Pt(16)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    p3 = doc.add_paragraph(
        "Para exames que são apenas digitalizados (imagens escaneadas), o uso de Reconhecimento Óptico de Caracteres (OCR) "
        "é indispensável. Abaixo listamos os principais algoritmos, frameworks e bibliotecas disponíveis no ecossistema "
        "Python localmente, ordenados do mais simples/leve ao mais avançado:"
    )
    p3.paragraph_format.space_after = Pt(10)
    
    # OCR List
    bullets = [
        ("RapidOCR", "Recomendado para implantação leve e CPU-only. Baseado em ONNX Runtime (com modelos extraídos do PaddleOCR), não possui dependências de binários externos no sistema operacional e roda inteiramente dentro do ambiente virtual Python de forma rápida."),
        ("Tesseract OCR (via pytesseract)", "O clássico da extração de texto. Requer a instalação do binário externo do Tesseract na máquina hospedeira e adição ao PATH. Excelente para texto corrido e linear, mas apresenta dificuldades significativas com equações matemáticas e colunas paralelas."),
        ("EasyOCR", "Construído sobre PyTorch e OpenCV, é altamente preciso e lida magistralmente com layouts complexos, mas exige mais hardware, download inicial de modelos profundos de ~1.5 GB e processamento substancialmente mais lento se rodar sem placa de vídeo dedicada (GPU).")
    ]
    
    for title, desc in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        run_title = bp.add_run(f"{title}: ")
        run_title.bold = True
        run_title.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        bp.add_run(desc)
        bp.paragraph_format.space_after = Pt(6)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Section 3: API & JSON Architecture
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Arquitetura do Pipeline e Formato de Saída (JSON)")
    h3_run.font.name = 'Segoe UI'
    h3_run.font.size = Pt(16)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    p4 = doc.add_paragraph(
        "O sistema está segmentado de forma modular, promovendo a separação de responsabilidades (SOLID). "
        "Os dados são ingeridos na pasta raiz e convertidos de forma estruturada para a pasta 'dados_processados' em dois tipos de saídas:"
    )
    p4.paragraph_format.space_after = Pt(10)
    
    sub1 = doc.add_paragraph()
    sub1.add_run("• JSON Individual por Questão: ").bold = True
    sub1.add_run("Salvo com o código único da questão (ex: FUVEST_2026_V1_Q11.json). Contém chaves explícitas para metadados, conteúdo textual (com campo de descrição para observações Note e Adote), matéria e lista de alternativas mapeando o gabarito oficial.")
    sub1.paragraph_format.space_after = Pt(6)
    
    sub2 = doc.add_paragraph()
    sub2.add_run("• JSON para Texto Complementar Comum: ").bold = True
    sub2.add_run("Salvo na pasta 'complementary'. Mapeia os enunciados compartilhados por múltiplas questões, indicando o array 'codigos_questoes' para cruzamento e montagem de visualizações dinâmicas no frontend.")
    sub2.paragraph_format.space_after = Pt(15)
    
    # Section 4: Study Roadmap Schedule (Table)
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Roteiro de Estudos Detalhado (4 Semanas)")
    h4_run.font.name = 'Segoe UI'
    h4_run.font.size = Pt(16)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    p5 = doc.add_paragraph(
        "Para dominar completamente a infraestrutura criada e avançar no desenvolvimento da API IFVest, "
        "siga o plano de estudo estruturado a seguir:"
    )
    p5.paragraph_format.space_after = Pt(12)
    
    # Create Table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Shading Accent 1'
    
    # Headers
    headers = ["Semana", "Módulo de Estudo", "Objetivo Prático"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.name = 'Segoe UI'
                r.font.size = Pt(10.5)
                
    # Table Content
    content = [
        ("Semana 1", "Python, Regex & Fundamentos do Pipeline", "Estudar expressôes regulares (re), depurar as strings geradas pelos PDFs e praticar a normalização de strings (Substituições Unicode em cleaners.py)."),
        ("Semana 2", "Modelagem com Pydantic e Ingestão Avançada", "Compreender tipos estritos do Pydantic (Field, BaseModel, Optional) e configurar validações dinâmicas para novos campos no arquivo schemas.py."),
        ("Semana 3", "Desenvolvimento REST com FastAPI & Uvicorn", "Aprofundar nos decoradores de rotas, injeção de dependências e expor parâmetros de filtros eficientes de busca no arquivo main.py."),
        ("Semana 4", "OCR Avançado, Imagens & Banco de Dados", "Praticar o processamento de imagens (recortar figuras usando coordenadas do pdfplumber) e planejar a persistência definitiva com SQLAlchemy / PostgreSQL.")
    ]
    
    for row_idx, (sem, mod, obj) in enumerate(content, start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = sem
        row_cells[1].text = mod
        row_cells[2].text = obj
        
        # Shade alternating rows
        fill = "F2F2F2" if row_idx % 2 == 0 else "FFFFFF"
        for cell in row_cells:
            set_cell_background(cell, fill)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Segoe UI'
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # Section 5: Instruction Manual
    h5 = doc.add_heading(level=1)
    h5_run = h5.add_run("5. Manual Prático: Execução e Testes")
    h5_run.font.name = 'Segoe UI'
    h5_run.font.size = Pt(16)
    h5_run.font.bold = True
    h5_run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    
    p6 = doc.add_paragraph(
        "Para validar localmente toda a suíte de ferramentas fornecida no seu projeto local, execute as seguintes "
        "instruções utilizando o terminal PowerShell ou CMD:"
    )
    p6.paragraph_format.space_after = Pt(10)
    
    cmds = [
        ("Passo A: Instalação das dependências", "pip install pypdf pdfplumber python-docx fastapi uvicorn pytest"),
        ("Passo B: Executar a Extração em Lote", "python run_extraction.py\n(Irá gerar 287 JSONs de questões limpas e validadas)"),
        ("Passo C: Inicializar o Servidor API", "uvicorn main:app --reload\n(Acesse a documentação interativa em http://127.0.0.1:8000/docs)"),
        ("Passo D: Validar e Consultar", "Utilizar rotas GET /questions ou GET /questions/{codigo} para testar os filtros de ano, edital e matéria.")
    ]
    
    for c_title, c_desc in cmds:
        cp = doc.add_paragraph()
        run_ctitle = cp.add_run(f"■ {c_title}:\n")
        run_ctitle.bold = True
        run_ctitle.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        run_cdesc = cp.add_run(c_desc)
        run_cdesc.italic = True
        cp.paragraph_format.left_indent = Inches(0.2)
        cp.paragraph_format.space_after = Pt(8)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(15)
    
    # Footer Note
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Fim do Roteiro. Desenvolvido para IFVest – Question Extraction API.")
    footer_run.font.size = Pt(9.5)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    # Save the file to the workspace root
    save_path = os.path.join(base_dir, "IFVest_Estudos_Roteiro.docx")
    doc.save(save_path)
    print(f"Document successfully created and saved to: {save_path}")

if __name__ == "__main__":
    create_roadmap()
